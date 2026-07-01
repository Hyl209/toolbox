from __future__ import annotations

import json
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SIDECAR = ROOT / "sidecar" / "hyl_sidecar.py"


def run_wordformatter(task: dict) -> list[dict]:
    with tempfile.TemporaryDirectory(prefix="hyl-sidecar-wordformatter-") as tmp:
        input_path = Path(tmp) / "input.json"
        input_path.write_text(json.dumps(task, ensure_ascii=False), encoding="utf-8")
        proc = subprocess.run(
            [sys.executable, str(SIDECAR), "run", "--tool", "wordformatter", "--input", str(input_path)],
            cwd=ROOT,
            text=True,
            encoding="utf-8",
            capture_output=True,
            check=False,
        )
    assert proc.stderr == ""
    lines = [line for line in proc.stdout.splitlines() if line]
    assert lines, "sidecar produced no stdout"
    return [json.loads(line) for line in lines]


def save_docx(path: Path, text: str = "body") -> Path:
    doc = Document()
    doc.add_paragraph(text)
    doc.save(str(path))
    return path


def test_wordformatter_default_config_exposes_page_and_styles() -> None:
    events = run_wordformatter({"task_id": "wordformatter-001", "action": "default_config", "payload": {}})

    event = events[-1]
    assert event["type"] == "result"
    assert "page" in event["data"]["config"]
    assert "body" in event["data"]["config"]["styles"]


def test_wordformatter_lists_docx_inputs_recursively() -> None:
    with tempfile.TemporaryDirectory(prefix="hyl-sidecar-wordformatter-files-") as tmp:
        root = Path(tmp)
        first = save_docx(root / "a.docx")
        (root / "~$a.docx").write_text("temp", encoding="utf-8")
        (root / "note.txt").write_text("ignore", encoding="utf-8")
        nested = root / "nested"
        nested.mkdir()
        save_docx(nested / "b.docx")

        events = run_wordformatter(
            {
                "task_id": "wordformatter-002",
                "action": "list",
                "payload": {"paths": [str(root), str(first)]},
            },
        )

    event = events[-1]
    assert event["type"] == "result"
    assert [item["name"] for item in event["data"]["files"]] == ["a.docx", "b.docx"]


def test_wordformatter_format_text_creates_docx() -> None:
    with tempfile.TemporaryDirectory(prefix="hyl-sidecar-wordformatter-files-") as tmp:
        root = Path(tmp)
        out_dir = root / "out"

        events = run_wordformatter(
            {
                "task_id": "wordformatter-003",
                "action": "format",
                "payload": {
                    "paths": [],
                    "text": "# Title\nbody text",
                    "output_dir": str(out_dir),
                    "output_mode": "copy",
                    "config": {"styles": {"heading1": {"size_pt": 20}, "body": {"size_pt": 11}}},
                },
            },
        )

        event = events[-1]
        assert event["type"] == "result"
        assert event["data"]["success_count"] == 1
        output = Path(event["data"]["results"][0]["output"])
        doc = Document(str(output))
        assert output.name == "word_text_formatted.docx"
        assert doc.paragraphs[0].text == "Title"
        assert doc.paragraphs[0].style.name == "Heading 1"
        assert doc.paragraphs[0].runs[0].font.size.pt == 20


def test_wordformatter_format_file_copy_uses_legacy_converter() -> None:
    with tempfile.TemporaryDirectory(prefix="hyl-sidecar-wordformatter-files-") as tmp:
        root = Path(tmp)
        source = save_docx(root / "source.docx", "body")
        out_dir = root / "out"

        events = run_wordformatter(
            {
                "task_id": "wordformatter-004",
                "action": "format",
                "payload": {
                    "paths": [str(source)],
                    "text": "",
                    "output_dir": str(out_dir),
                    "output_mode": "copy",
                    "config": {"styles": {"body": {"size_pt": 13}}},
                },
            },
        )

        event = events[-1]
        assert event["type"] == "result"
        output = Path(event["data"]["results"][0]["output"])
        assert output.name == "source_formatted.docx"
        assert source.exists()
        formatted = Document(str(output))
        assert formatted.paragraphs[0].runs[0].font.size.pt == 13


def test_wordformatter_rejects_missing_output_dir_for_copy() -> None:
    events = run_wordformatter(
        {
            "task_id": "wordformatter-005",
            "action": "format",
            "payload": {"text": "body", "output_mode": "copy"},
        },
    )

    event = events[-1]
    assert event["type"] == "error"
    assert event["code"] == "INVALID_PAYLOAD"
