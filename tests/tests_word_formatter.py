from __future__ import annotations

import importlib.util
from pathlib import Path

import pytest
from docx import Document


ROOT = Path(__file__).resolve().parents[1]
MODULE_PATH = ROOT / 'modules' / 'word-formatter' / 'converter.py'


def load_converter():
    spec = importlib.util.spec_from_file_location('tests_word_formatter_converter', MODULE_PATH)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(module)
    return module


def save_demo_doc(path: Path, text: str = 'hello') -> Path:
    doc = Document()
    doc.add_paragraph(text)
    doc.save(str(path))
    return path


def test_collect_word_inputs_filters_docx_recursively_and_skips_temp_files(tmp_path):
    module = load_converter()
    root = tmp_path / 'root'
    root.mkdir()
    first = save_demo_doc(root / 'a.docx')
    (root / '~$a.docx').write_text('temp', encoding='utf-8')
    (root / 'note.txt').write_text('x', encoding='utf-8')
    nested = root / 'nested'
    nested.mkdir()
    second = save_demo_doc(nested / 'b.docx')

    result = module.collect_word_inputs([str(root), str(first)])

    assert result == sorted([first.resolve(), second.resolve()])


def test_parse_markdown_text_detects_heading_levels_and_body():
    module = load_converter()

    items = module.parse_markdown_text('# Title\nbody\n### Deep')

    assert items == [
        {'type': 'heading', 'level': 1, 'text': 'Title'},
        {'type': 'body', 'level': 0, 'text': 'body'},
        {'type': 'heading', 'level': 3, 'text': 'Deep'},
    ]


def test_create_docx_from_text_applies_heading_and_body_styles(tmp_path):
    module = load_converter()
    config = module.get_default_config()
    config['styles']['heading1']['font'] = 'Arial'
    config['styles']['heading1']['size_pt'] = 20
    config['styles']['body']['font'] = 'Arial'
    config['styles']['body']['size_pt'] = 11

    output = module.create_docx_from_text('# Title\nbody text', config, tmp_path)

    doc = Document(str(output))
    assert output.name == 'word_text_formatted.docx'
    assert doc.paragraphs[0].style.name == 'Heading 1'
    assert doc.paragraphs[0].text == 'Title'
    assert doc.paragraphs[0].runs[0].font.size.pt == 20
    assert doc.paragraphs[1].style.name == 'Normal'
    assert doc.paragraphs[1].runs[0].font.size.pt == 11


def test_format_docx_file_applies_page_heading_body_and_table_styles(tmp_path):
    module = load_converter()
    source = tmp_path / 'source.docx'
    doc = Document()
    doc.add_paragraph('# Markdown Title')
    doc.add_paragraph('body')
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = 'cell'
    doc.save(str(source))
    config = module.get_default_config()
    config['page']['top_margin_cm'] = 2
    config['styles']['heading1']['font'] = 'Arial'
    config['styles']['heading1']['size_pt'] = 18
    config['styles']['body']['size_pt'] = 12
    config['styles']['table']['size_pt'] = 9

    output = module.format_docx_file(source, config, tmp_path, 'copy')

    formatted = Document(str(output))
    assert output.name == 'source_formatted.docx'
    assert round(formatted.sections[0].top_margin.cm, 1) == 2.0
    assert formatted.paragraphs[0].text == 'Markdown Title'
    assert formatted.paragraphs[0].style.name == 'Heading 1'
    assert formatted.paragraphs[0].runs[0].font.size.pt == 18
    assert formatted.paragraphs[1].runs[0].font.size.pt == 12
    assert formatted.tables[0].cell(0, 0).paragraphs[0].runs[0].font.size.pt == 9


def test_format_docx_file_overwrite_writes_back_to_source(tmp_path):
    module = load_converter()
    source = save_demo_doc(tmp_path / 'source.docx', 'body')
    config = module.get_default_config()
    config['styles']['body']['size_pt'] = 13

    output = module.format_docx_file(source, config, tmp_path, 'overwrite')

    assert output == source.resolve()
    formatted = Document(str(source))
    assert formatted.paragraphs[0].runs[0].font.size.pt == 13


def test_format_docx_file_overwrite_preserves_source_when_save_fails(tmp_path, monkeypatch):
    module = load_converter()
    source = tmp_path / 'source.docx'
    source.write_bytes(b'original')

    class BrokenDocument:
        def save(self, path: str):
            Path(path).write_bytes(b'partial')
            raise RuntimeError('save failed')

    monkeypatch.setattr(module, 'Document', lambda _: BrokenDocument())
    monkeypatch.setattr(module, 'apply_document_format', lambda _document, _config: None)

    with pytest.raises(RuntimeError, match='save failed'):
        module.format_docx_file(source, module.get_default_config(), '', 'overwrite')

    assert source.read_bytes() == b'original'


def test_validate_request_rejects_empty_input_text_overwrite_and_missing_output(tmp_path):
    module = load_converter()
    source = save_demo_doc(tmp_path / 'source.docx', 'body')

    assert '请拖入 Word 文件或输入文本' in module.validate_request([], '', '', 'copy')
    assert '直接文本输入不能使用原地覆盖' in module.validate_request([], 'body', str(tmp_path), 'overwrite')
    assert '请选择输出目录' in module.validate_request([source], '', '', 'copy')
    assert module.validate_request([source], '', '', 'overwrite') == []


def test_validate_config_reports_invalid_numbers():
    module = load_converter()
    config = module.get_default_config()
    config['page']['top_margin_cm'] = 'bad'
    config['styles']['body']['size_pt'] = '0'

    errors = module.validate_config(config)

    assert any('top_margin_cm' in item for item in errors)
    assert any('body.size_pt' in item for item in errors)
