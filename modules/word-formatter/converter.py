from __future__ import annotations

from copy import deepcopy
import os
from pathlib import Path
import re
import tempfile

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
except ModuleNotFoundError:  # pragma: no cover - validated at runtime
    Document = None
    qn = None
    Cm = None
    Pt = None


SUPPORTED_INPUT_EXTENSIONS = {'.docx'}
OUTPUT_MODE_COPY = 'copy'
OUTPUT_MODE_OVERWRITE = 'overwrite'
STYLE_KEYS = ('heading1', 'heading2', 'heading3', 'heading4', 'body', 'table')

DEFAULT_PAGE_CONFIG = {
    'top_margin_cm': 2.54,
    'bottom_margin_cm': 2.54,
    'left_margin_cm': 3.18,
    'right_margin_cm': 3.18,
    'header_distance_cm': 1.5,
    'footer_distance_cm': 1.75,
}

DEFAULT_STYLE_CONFIG = {
    'heading1': {'font': 'Microsoft YaHei', 'size_pt': 18.0, 'bold': True, 'line_spacing': 1.5, 'space_before_pt': 12.0, 'space_after_pt': 6.0, 'first_line_indent_cm': 0.0},
    'heading2': {'font': 'Microsoft YaHei', 'size_pt': 16.0, 'bold': True, 'line_spacing': 1.5, 'space_before_pt': 10.0, 'space_after_pt': 6.0, 'first_line_indent_cm': 0.0},
    'heading3': {'font': 'Microsoft YaHei', 'size_pt': 14.0, 'bold': True, 'line_spacing': 1.5, 'space_before_pt': 8.0, 'space_after_pt': 4.0, 'first_line_indent_cm': 0.0},
    'heading4': {'font': 'Microsoft YaHei', 'size_pt': 12.0, 'bold': True, 'line_spacing': 1.5, 'space_before_pt': 6.0, 'space_after_pt': 4.0, 'first_line_indent_cm': 0.0},
    'body': {'font': 'Microsoft YaHei', 'size_pt': 12.0, 'bold': False, 'line_spacing': 1.5, 'space_before_pt': 0.0, 'space_after_pt': 0.0, 'first_line_indent_cm': 0.74},
    'table': {'font': 'Microsoft YaHei', 'size_pt': 10.5, 'bold': False, 'line_spacing': 1.2, 'space_before_pt': 0.0, 'space_after_pt': 0.0, 'first_line_indent_cm': 0.0},
}

DEFAULT_CONFIG = {
    'page': DEFAULT_PAGE_CONFIG,
    'styles': DEFAULT_STYLE_CONFIG,
}

_MARKDOWN_HEADING_RE = re.compile(r'^(#{1,4})\s+(.+?)\s*$')


class WordFormatError(Exception):
    pass


def get_default_config() -> dict[str, object]:
    return deepcopy(DEFAULT_CONFIG)


def collect_word_inputs(paths: list[str]) -> list[Path]:
    unique: dict[Path, None] = {}
    for raw in paths:
        path = Path(raw).resolve()
        try:
            if path.is_file() and _is_supported_docx(path):
                unique[path] = None
            elif path.is_dir():
                for item in sorted(path.rglob('*')):
                    if item.is_file() and _is_supported_docx(item):
                        unique[item.resolve()] = None
        except OSError:
            continue
    return sorted(unique.keys())


def _is_supported_docx(path: Path) -> bool:
    return path.suffix.lower() in SUPPORTED_INPUT_EXTENSIONS and not path.name.startswith('~$')


def parse_markdown_text(text: str) -> list[dict[str, object]]:
    items: list[dict[str, object]] = []
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        match = _MARKDOWN_HEADING_RE.match(line)
        if match:
            items.append({'type': 'heading', 'level': len(match.group(1)), 'text': match.group(2).strip()})
        else:
            items.append({'type': 'body', 'level': 0, 'text': line})
    return items


def validate_request(files: list[Path], text: str, output_dir: str, output_mode: str) -> list[str]:
    errors: list[str] = []
    has_files = bool(files)
    has_text = bool(text.strip())
    if not has_files and not has_text:
        errors.append('请拖入 Word 文件或输入文本')
    if has_text and output_mode == OUTPUT_MODE_OVERWRITE:
        errors.append('直接文本输入不能使用原地覆盖')
    if output_mode not in {OUTPUT_MODE_COPY, OUTPUT_MODE_OVERWRITE}:
        errors.append('输出模式不正确')
    if (output_mode == OUTPUT_MODE_COPY or has_text) and not output_dir.strip():
        errors.append('请选择输出目录')
    return errors


def validate_config(config: dict[str, object]) -> list[str]:
    errors: list[str] = []
    page = dict(config.get('page', {}))
    styles = dict(config.get('styles', {}))
    for key in DEFAULT_PAGE_CONFIG:
        _append_positive_number_error(errors, page.get(key), key)
    for style_key in STYLE_KEYS:
        style = dict(styles.get(style_key, {}))
        for key in ('size_pt', 'line_spacing'):
            _append_positive_number_error(errors, style.get(key), f'{style_key}.{key}')
        for key in ('space_before_pt', 'space_after_pt', 'first_line_indent_cm'):
            _append_non_negative_number_error(errors, style.get(key), f'{style_key}.{key}')
    return errors


def _append_positive_number_error(errors: list[str], value: object, field: str) -> None:
    try:
        if float(value) <= 0:
            errors.append(f'{field} 必须大于 0')
    except (TypeError, ValueError):
        errors.append(f'{field} 必须是数字')


def _append_non_negative_number_error(errors: list[str], value: object, field: str) -> None:
    try:
        if float(value) < 0:
            errors.append(f'{field} 不能小于 0')
    except (TypeError, ValueError):
        errors.append(f'{field} 必须是数字')


def format_docx_file(input_path: str | Path, config: dict[str, object], output_dir: str | Path = '', output_mode: str = OUTPUT_MODE_COPY) -> Path:
    if Document is None:
        raise WordFormatError('未安装 python-docx，无法处理 Word 文档')
    source = Path(input_path).resolve()
    if not _is_supported_docx(source):
        raise WordFormatError('仅支持 .docx 文件')
    errors = validate_config(config)
    if errors:
        raise WordFormatError('\n'.join(errors))
    document = Document(str(source))
    apply_document_format(document, config)
    if output_mode == OUTPUT_MODE_OVERWRITE:
        output_path = source
    else:
        output_path = _build_copy_output_path(source, Path(output_dir))
    output_path.parent.mkdir(parents=True, exist_ok=True)
    if output_mode == OUTPUT_MODE_OVERWRITE:
        _save_document_atomic(document, output_path)
    else:
        document.save(str(output_path))
    return output_path


def _save_document_atomic(document, output_path: Path) -> None:
    tmp_path = ''
    try:
        with tempfile.NamedTemporaryFile(
            dir=output_path.parent,
            prefix=f'.{output_path.stem}.',
            suffix=output_path.suffix,
            delete=False,
        ) as tmp:
            tmp_path = tmp.name
        document.save(tmp_path)
        os.replace(tmp_path, output_path)
        tmp_path = ''
    finally:
        if tmp_path:
            try:
                Path(tmp_path).unlink()
            except OSError:
                pass


def create_docx_from_text(text: str, config: dict[str, object], output_dir: str | Path, output_name: str = 'word_text_formatted.docx') -> Path:
    if Document is None:
        raise WordFormatError('未安装 python-docx，无法生成 Word 文档')
    items = parse_markdown_text(text)
    if not items:
        raise WordFormatError('请输入文本内容')
    errors = validate_config(config)
    if errors:
        raise WordFormatError('\n'.join(errors))
    document = Document()
    for item in items:
        content = str(item['text'])
        if item['type'] == 'heading':
            document.add_heading(content, level=int(item['level']))
        else:
            document.add_paragraph(content)
    apply_document_format(document, config)
    output_path = _resolve_name_conflict(Path(output_dir) / output_name)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    return output_path


def format_batch(files: list[Path], text: str, config: dict[str, object], output_dir: str, output_mode: str) -> list[Path]:
    errors = validate_request(files, text, output_dir, output_mode)
    if errors:
        raise WordFormatError('\n'.join(errors))
    outputs: list[Path] = []
    for file in files:
        outputs.append(format_docx_file(file, config, output_dir, output_mode))
    if text.strip():
        outputs.append(create_docx_from_text(text, config, output_dir))
    return outputs


def apply_document_format(document, config: dict[str, object]) -> None:
    page = _merged_page_config(config)
    styles = _merged_style_config(config)
    _apply_page_config(document, page)
    _apply_named_styles(document, styles)
    for paragraph in document.paragraphs:
        style_key = _normalize_paragraph_style(paragraph)
        _apply_paragraph_format(paragraph, styles[style_key])
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.style = document.styles['Normal']
                    _apply_paragraph_format(paragraph, styles['table'])


def _merged_page_config(config: dict[str, object]) -> dict[str, float]:
    page = dict(DEFAULT_PAGE_CONFIG)
    page.update(dict(config.get('page', {})))
    try:
        return {key: float(value) for key, value in page.items()}
    except (ValueError, TypeError) as exc:
        raise WordFormatError(f'页面配置包含无效数值: {exc}') from exc


def _merged_style_config(config: dict[str, object]) -> dict[str, dict[str, object]]:
    configured = dict(config.get('styles', {}))
    styles: dict[str, dict[str, object]] = {}
    for key in STYLE_KEYS:
        merged = dict(DEFAULT_STYLE_CONFIG[key])
        merged.update(dict(configured.get(key, {})))
        styles[key] = merged
    return styles


def _apply_page_config(document, page: dict[str, float]) -> None:
    for section in document.sections:
        section.top_margin = Cm(page['top_margin_cm'])
        section.bottom_margin = Cm(page['bottom_margin_cm'])
        section.left_margin = Cm(page['left_margin_cm'])
        section.right_margin = Cm(page['right_margin_cm'])
        section.header_distance = Cm(page['header_distance_cm'])
        section.footer_distance = Cm(page['footer_distance_cm'])


def _apply_named_styles(document, styles: dict[str, dict[str, object]]) -> None:
    style_names = {
        'heading1': 'Heading 1',
        'heading2': 'Heading 2',
        'heading3': 'Heading 3',
        'heading4': 'Heading 4',
        'body': 'Normal',
    }
    for key, name in style_names.items():
        try:
            style = document.styles[name]
        except KeyError:
            continue
        _apply_font_to_style(style, styles[key])
        _apply_paragraph_format_to_style(style, styles[key])


def _normalize_paragraph_style(paragraph) -> str:
    text = paragraph.text.strip()
    match = _MARKDOWN_HEADING_RE.match(text)
    if match:
        level = len(match.group(1))
        paragraph.text = match.group(2).strip()
        paragraph.style = f'Heading {level}'
        return f'heading{level}'
    style_name = getattr(paragraph.style, 'name', '')
    if style_name in {'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4'}:
        return f'heading{style_name[-1]}'
    paragraph.style = 'Normal'
    return 'body'


def _apply_font_to_style(style, style_config: dict[str, object]) -> None:
    font = style.font
    font.name = str(style_config['font'])
    font.size = Pt(float(style_config['size_pt']))
    font.bold = bool(style_config['bold'])
    if qn is not None and font.element.rPr is not None:
        font.element.rPr.rFonts.set(qn('w:eastAsia'), str(style_config['font']))


def _apply_paragraph_format_to_style(style, style_config: dict[str, object]) -> None:
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = float(style_config['line_spacing'])
    paragraph_format.space_before = Pt(float(style_config['space_before_pt']))
    paragraph_format.space_after = Pt(float(style_config['space_after_pt']))
    paragraph_format.first_line_indent = Cm(float(style_config['first_line_indent_cm']))


def _apply_paragraph_format(paragraph, style_config: dict[str, object]) -> None:
    _apply_paragraph_format_to_style(paragraph, style_config)
    for run in paragraph.runs:
        run.font.name = str(style_config['font'])
        run.font.size = Pt(float(style_config['size_pt']))
        run.font.bold = bool(style_config['bold'])
        if qn is not None and run.font.element.rPr is not None:
            run.font.element.rPr.rFonts.set(qn('w:eastAsia'), str(style_config['font']))


def _build_copy_output_path(source: Path, output_dir: Path) -> Path:
    return _resolve_name_conflict(output_dir / f'{source.stem}_formatted.docx')


def _resolve_name_conflict(path: Path) -> Path:
    if not path.exists():
        return path
    for index in range(1, 1000):
        candidate = path.with_name(f'{path.stem}({index}){path.suffix}')
        if not candidate.exists():
            return candidate
    raise WordFormatError('无法生成不冲突的输出文件名')
